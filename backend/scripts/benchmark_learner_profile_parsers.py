"""
@Date: 2026-05-25
@Author: xisy
@Discription: 学情 docx 解析方案对比基准
    对比 mammoth / python-docx / pandoc 三种本地解析方式的速度与抽取效果，
    用于评估是否可替换当前 MinerU 远程批处理链路
"""

import io
import subprocess
import time
from pathlib import Path

import mammoth
from docx import Document

from app.modules.learner_profile.rules import parse_learner_profile_text


PROFILE_DIR = Path("/Users/xisy/Personal/Projects/EduWeave/教育赛题/学情分析")


def load_profile_files() -> list[tuple[str, bytes]]:
    """加载所有非临时学情 docx 文件。"""
    items: list[tuple[str, bytes]] = []
    for path in sorted(PROFILE_DIR.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        items.append((path.name, path.read_bytes()))
    return items


def parse_with_mammoth(content: bytes) -> str:
    """使用 mammoth 转换 docx 为 markdown。"""
    result = mammoth.convert_to_markdown(io.BytesIO(content))
    return result.value


def parse_with_python_docx(content: bytes) -> str:
    """使用 python-docx 抽取段落与表格构造 markdown。"""
    document = Document(io.BytesIO(content))
    lines: list[str] = []
    body_elements = document.element.body
    table_iter = iter(document.tables)
    paragraph_iter = iter(document.paragraphs)
    paragraph_map = {paragraph._element: paragraph for paragraph in document.paragraphs}
    table_map = {table._element: table for table in document.tables}
    for child in body_elements.iterchildren():
        if child in paragraph_map:
            paragraph = paragraph_map[child]
            text_value = paragraph.text.strip()
            if text_value:
                lines.append(text_value)
        elif child in table_map:
            table = table_map[child]
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        lines.append(cell_text)
    return "\n".join(lines)


def parse_with_pandoc(content: bytes) -> str:
    """使用 pandoc CLI 转换 docx 为 markdown。"""
    process_result = subprocess.run(
        ["pandoc", "-f", "docx", "-t", "gfm", "--wrap=none"],
        input=content,
        capture_output=True,
        check=True,
        timeout=30,
    )
    return process_result.stdout.decode("utf-8", errors="ignore")


def benchmark_method(method_name: str, parser_callable, files: list[tuple[str, bytes]]) -> dict:
    """对单一解析方案跑全量并统计耗时与抽取记录数。"""
    per_file_durations: list[float] = []
    extract_counts: list[int] = []
    sample_markdown: str = ""
    sample_filename: str = ""
    total_start = time.perf_counter()
    for filename, content in files:
        single_start = time.perf_counter()
        markdown_text = parser_callable(content)
        single_elapsed_ms = (time.perf_counter() - single_start) * 1000
        per_file_durations.append(single_elapsed_ms)
        parse_result = parse_learner_profile_text(
            markdown_text,
            fallback_title=filename,
            fallback_filename=filename,
        )
        extract_counts.append(len(parse_result.records))
        if not sample_markdown:
            sample_markdown = markdown_text
            sample_filename = filename
    total_elapsed_ms = (time.perf_counter() - total_start) * 1000
    return {
        "method_name": method_name,
        "total_ms": total_elapsed_ms,
        "avg_ms": sum(per_file_durations) / len(per_file_durations) if per_file_durations else 0.0,
        "min_ms": min(per_file_durations) if per_file_durations else 0.0,
        "max_ms": max(per_file_durations) if per_file_durations else 0.0,
        "records_total": sum(extract_counts),
        "records_per_file": extract_counts,
        "sample_filename": sample_filename,
        "sample_markdown_head": sample_markdown[:600],
    }


def main() -> None:
    files = load_profile_files()
    print(f"待测试学情文件数: {len(files)}\n")

    methods = [
        ("mammoth", parse_with_mammoth),
        ("python-docx", parse_with_python_docx),
        ("pandoc", parse_with_pandoc),
    ]
    benchmark_results: list[dict] = []
    for method_name, parser_callable in methods:
        try:
            stats = benchmark_method(method_name, parser_callable, files)
        except Exception as exc:  # noqa: BLE001
            print(f"[{method_name}] 失败: {exc}\n")
            continue
        benchmark_results.append(stats)
        print(f"== {method_name} ==")
        print(f"  total: {stats['total_ms']:.1f} ms  | avg: {stats['avg_ms']:.1f} ms  | min: {stats['min_ms']:.1f} ms  | max: {stats['max_ms']:.1f} ms")
        print(f"  records 累计: {stats['records_total']} | 每文件: {stats['records_per_file']}")
        print(f"  样例 [{stats['sample_filename']}] markdown 前 600 字符：")
        print("  " + stats["sample_markdown_head"].replace("\n", "\n  "))
        print()

    if benchmark_results:
        print("== 汇总 ==")
        for stats in benchmark_results:
            print(
                f"  {stats['method_name']:<12} avg={stats['avg_ms']:.1f}ms total={stats['total_ms']:.1f}ms records={stats['records_total']}"
            )


if __name__ == "__main__":
    main()
