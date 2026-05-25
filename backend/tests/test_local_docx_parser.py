"""
@Date: 2026-05-25
@Author: xisy
@Discription: 本地 docx 解析服务与学情规则衔接测试
"""

from io import BytesIO

import pytest
from docx import Document

from app.core.exceptions import AppException, BusinessErrorCode
from app.modules.learner_profile.rules import parse_learner_profile_text
from app.shared.document.docx_parser import LOCAL_DOCX_MODEL_VERSION, LocalDocxParseService


# 在 conftest autouse 桩替换前抓取真实方法引用，便于本测试模块按需还原
_REAL_PARSE_DOCUMENT = LocalDocxParseService.parse_document


@pytest.fixture(autouse=True)
def restore_real_local_docx_parser(monkeypatch: pytest.MonkeyPatch):
    """本测试需要走真实 LocalDocxParseService，还原被 conftest 桩替换的方法。"""
    monkeypatch.setattr(LocalDocxParseService, "parse_document", _REAL_PARSE_DOCUMENT)


def build_minimal_learner_profile_docx() -> bytes:
    """用 python-docx 构造一份与赛题样例同构的最小学情 docx。"""
    document = Document()
    document.add_paragraph("张测试 — 学情分析")
    document.add_paragraph("一、基本信息")
    info_table = document.add_table(rows=4, cols=2)
    info_pairs = [
        ("姓名", "张测试"),
        ("所属地区", "北京"),
        ("年级", "三年级"),
        ("学习科目", "语文、数学"),
    ]
    for row_index, (label, value) in enumerate(info_pairs):
        info_table.rows[row_index].cells[0].text = label
        info_table.rows[row_index].cells[1].text = value
    document.add_paragraph("二、使用教材")
    textbook_table = document.add_table(rows=3, cols=2)
    textbook_table.rows[0].cells[0].text = "科目"
    textbook_table.rows[0].cells[1].text = "教材版本"
    textbook_table.rows[1].cells[0].text = "语文"
    textbook_table.rows[1].cells[1].text = "人民教育出版社-语文-三年级下册"
    textbook_table.rows[2].cells[0].text = "数学"
    textbook_table.rows[2].cells[1].text = "北京出版社-数学-三年级下册"
    document.add_paragraph("三、科目成绩")
    document.add_paragraph("语文：92分（月考）")
    document.add_paragraph("数学：88分（期末质检）")
    document.add_paragraph("四、学生基本情况描述")
    document.add_paragraph(
        "该学生阅读理解能力较强，作文思路清晰；数学方面基础计算能力较好，"
        "但乘法口诀记忆尚不牢固，遇到需要逆向思维的题目时容易卡壳。"
    )
    document.add_paragraph("五、培训时间规划")
    document.add_paragraph(
        "计划每周安排2次数学课（每次2课时），重点训练乘法运算能力；"
        "每周1次语文课（每次2课时），侧重阅读理解能力。"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parse_document_outputs_expected_structure():
    """真实 docx 经本地解析后产物结构符合下游约定。"""
    docx_bytes = build_minimal_learner_profile_docx()
    parser = LocalDocxParseService()
    normalized_document = parser.parse_document(
        file_name="张测试.docx",
        content=docx_bytes,
        data_id="profile_test_task_1",
    )

    assert normalized_document.batch_id == "local-profile_test_task_1"
    assert normalized_document.model_version == LOCAL_DOCX_MODEL_VERSION
    for section_title in ("一、基本信息", "二、使用教材", "三、科目成绩", "四、学生基本情况描述", "五、培训时间规划"):
        assert section_title in normalized_document.markdown_text
    assert normalized_document.full_zip_bytes == docx_bytes
    assert normalized_document.asset_files == {}
    assert len(normalized_document.pages) == 1
    assert normalized_document.pages[0].page_no == 1
    assert any(item.get("type") == "table_cell" for item in normalized_document.content_list_json)


def test_parse_document_feeds_learner_profile_rules():
    """解析产物可被现有学情规则正确切分出多学科记录。"""
    docx_bytes = build_minimal_learner_profile_docx()
    parser = LocalDocxParseService()
    normalized_document = parser.parse_document(
        file_name="张测试.docx",
        content=docx_bytes,
        data_id="profile_test_task_2",
    )

    parse_result = parse_learner_profile_text(
        normalized_document.markdown_text,
        fallback_title="张测试学情",
        fallback_filename="张测试.docx",
    )

    assert parse_result.student_name == "张测试"
    assert parse_result.grade_code == "grade_3"
    subject_codes = {record.subject_code for record in parse_result.records}
    assert {"chinese", "math"} <= subject_codes
    score_by_subject = {record.subject_code: record.score_value for record in parse_result.records}
    assert score_by_subject["chinese"] == 92.0
    assert score_by_subject["math"] == 88.0


def test_parse_document_rejects_empty_content():
    """空 docx 字节直接报错，不向下游抛出脏数据。"""
    parser = LocalDocxParseService()
    with pytest.raises(AppException) as excinfo:
        parser.parse_document(file_name="empty.docx", content=b"", data_id="profile_test_task_3")
    assert excinfo.value.code == BusinessErrorCode.INVALID_FILE_TYPE


def test_parse_document_rejects_invalid_docx_bytes():
    """无效 docx 字节直接报错，便于上层任务统一处理失败。"""
    parser = LocalDocxParseService()
    with pytest.raises(AppException) as excinfo:
        parser.parse_document(
            file_name="invalid.docx",
            content=b"not-a-real-docx",
            data_id="profile_test_task_4",
        )
    assert excinfo.value.code == BusinessErrorCode.INVALID_FILE_TYPE
