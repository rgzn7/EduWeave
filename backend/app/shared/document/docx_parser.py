"""
@Date: 2026-05-25
@Author: xisy
@Discription: 本地 docx 解析服务，输出归一化 NormalizedDocument 供下游复用
"""

from io import BytesIO
from pathlib import PurePosixPath
from typing import Any

from docx import Document
from docx.parts.image import ImagePart

from app.core.exceptions import AppException, BusinessErrorCode
from app.shared.mineru.schemas import NormalizedBlock, NormalizedDocument, NormalizedPage


# 本地 docx 解析模型版本标识，便于和 MinerU 产物区分
LOCAL_DOCX_MODEL_VERSION = "local_docx"
# docx 段落与图片资源在 content_list 中的页索引（docx 无明确页概念，统一归到第 1 页）
_DEFAULT_PAGE_IDX = 0
_ASSET_ROOT_DIR = "assets"


class LocalDocxParseService:
    """基于 python-docx 的本地 docx 解析服务。"""

    def parse_document(
        self,
        *,
        file_name: str,
        content: bytes,
        data_id: str,
    ) -> NormalizedDocument:
        """解析 docx 字节流并产出归一化文档结构。"""
        if not content:
            raise AppException(BusinessErrorCode.INVALID_FILE_TYPE, "学情文件为空", {"file_name": file_name})
        try:
            document = Document(BytesIO(content))
        except Exception as exc:  # noqa: BLE001
            raise AppException(
                BusinessErrorCode.INVALID_FILE_TYPE,
                "学情文件无法解析为 docx",
                {"file_name": file_name, "error": str(exc)},
            ) from exc

        markdown_lines, content_list_json = self._extract_body_blocks(document)
        asset_files = self._extract_image_assets(document)
        for relative_path in asset_files:
            content_list_json.append(
                {
                    "page_idx": _DEFAULT_PAGE_IDX,
                    "type": "image",
                    "img_path": relative_path,
                }
            )
        markdown_text = "\n".join(markdown_lines).strip()
        pages = self._build_single_page(markdown_text, content_list_json)
        return NormalizedDocument(
            batch_id=f"local-{data_id}",
            file_name=file_name,
            data_id=data_id,
            model_version=LOCAL_DOCX_MODEL_VERSION,
            markdown_text=markdown_text,
            content_list_json=content_list_json,
            pages=pages,
            # docx 本身就是 zip 容器，直接归档原始字节即可保留完整源
            full_zip_bytes=content,
            asset_files=asset_files,
            raw_metadata={
                "parser": LOCAL_DOCX_MODEL_VERSION,
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "asset_count": len(asset_files),
            },
        )

    def _extract_body_blocks(self, document: Document) -> tuple[list[str], list[dict[str, Any]]]:
        """按 body 元素顺序抽取段落与表格文本。"""
        paragraph_map = {paragraph._element: paragraph for paragraph in document.paragraphs}
        table_map = {table._element: table for table in document.tables}
        markdown_lines: list[str] = []
        content_list_json: list[dict[str, Any]] = []
        for child in document.element.body.iterchildren():
            if child in paragraph_map:
                text_value = paragraph_map[child].text.strip()
                if not text_value:
                    continue
                markdown_lines.append(text_value)
                content_list_json.append(
                    {
                        "page_idx": _DEFAULT_PAGE_IDX,
                        "type": "paragraph",
                        "text": text_value,
                    }
                )
            elif child in table_map:
                for row in table_map[child].rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if not cell_text:
                            continue
                        markdown_lines.append(cell_text)
                        content_list_json.append(
                            {
                                "page_idx": _DEFAULT_PAGE_IDX,
                                "type": "table_cell",
                                "text": cell_text,
                            }
                        )
        return markdown_lines, content_list_json

    def _extract_image_assets(self, document: Document) -> dict[str, bytes]:
        """抽取 docx 内嵌图片，按统一命名落到 assets/ 目录下。"""
        asset_files: dict[str, bytes] = {}
        for related_part in document.part.related_parts.values():
            if not isinstance(related_part, ImagePart):
                continue
            partname = PurePosixPath(str(related_part.partname))
            relative_path = f"{_ASSET_ROOT_DIR}/{partname.name}"
            # 同名图片去重，保留首个
            if relative_path in asset_files:
                continue
            asset_files[relative_path] = related_part.blob
        return asset_files

    @staticmethod
    def _build_single_page(markdown_text: str, content_list_json: list[dict[str, Any]]) -> list[NormalizedPage]:
        """docx 无页码概念，统一归到第 1 页。"""
        if not markdown_text and not content_list_json:
            return []
        block = NormalizedBlock(
            page_no=1,
            block_no=1,
            block_type="paragraph",
            text_content=markdown_text or None,
            markdown_content=markdown_text or None,
        )
        return [
            NormalizedPage(
                page_no=1,
                text_content=markdown_text or None,
                markdown_content=markdown_text or None,
                layout_json={"raw_items": content_list_json},
                blocks=[block],
            )
        ]

