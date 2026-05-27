"""
@Date: 2026-05-09
@Author: xisy
@Discription: DOCX 渲染与导出归档服务
"""

import hashlib
import json
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.exceptions import AppException, BusinessErrorCode
from app.modules.file_asset.schemas import FileDownloadUrlResponse
from app.modules.p0_models import FileObject
from app.shared.document.labels import (
    DIFFICULTY_LEVEL_LABELS,
    LESSON_AFTER_CLASS_LABELS,
    LESSON_COURSE_OVERVIEW_LABELS,
    QUESTION_TYPE_LABELS,
    SCENE_TYPE_LABELS,
    iter_known_fields,
    labelize,
)
from app.shared.document.naming import strip_lesson_prefix
from app.shared.storage import ObsStorageClient
from app.shared.utils import DateTimeUtil

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# 模板版本号：DOCX 渲染规则发生破坏性变化时 +1，object_key 会嵌入 tv{N} 段，
# 旧 export_file_id 自然失效，前端再次触发 /export-docx 后会落到新 key。
DOCX_TEMPLATE_VERSION = 2


class DocxRenderService:
    """根据结构化结果渲染 DOCX 文档。"""

    def render_curriculum_plan(self, plan) -> bytes:
        """渲染课程大纲 DOCX。"""
        content = _ensure_dict(plan.content_json)
        document = _create_document()
        document.add_heading(_safe_text(plan.plan_title, "课程大纲"), level=0)
        _add_meta_paragraph(document, "课程摘要", plan.summary_text)
        _add_meta_paragraph(document, "总课次", plan.course_count)
        _add_meta_paragraph(document, "单次课时", f"{plan.session_duration_minutes} 分钟")

        _add_labeled_section(document, "课程概览", content.get("course_overview"), LESSON_COURSE_OVERVIEW_LABELS)
        _add_list_section(document, "阶段目标", content.get("stage_goals"))
        _add_list_section(document, "课程重点", content.get("key_points"))
        _add_list_section(document, "课程难点", content.get("difficult_points"))
        _add_list_section(document, "学情适配策略", content.get("learner_adjustments"))
        _add_lesson_session_table(document, content.get("lesson_sessions"))
        return _dump_document(document)

    def render_lesson_plan(
        self,
        lesson_plan,
        *,
        knowledge_point_names: dict[int, str] | None = None,
    ) -> bytes:
        """渲染教案 DOCX。knowledge_point_names 用于把 id 列表替换成知识点名称。"""
        content = _ensure_dict(lesson_plan.content_json)
        kp_names = knowledge_point_names or {}
        document = _create_document()
        document.add_heading(strip_lesson_prefix(lesson_plan.lesson_title) or "教案", level=0)
        _add_meta_paragraph(document, "教案摘要", lesson_plan.summary_text)
        if lesson_plan.class_session_no is not None:
            _add_meta_paragraph(document, "课次", f"第 {lesson_plan.class_session_no} 讲")

        _add_labeled_section(document, "课程概述", content.get("course_overview"), LESSON_COURSE_OVERVIEW_LABELS)
        _add_list_section(document, "物料清单", content.get("material_list"))
        _add_list_section(document, "核心知识", content.get("core_knowledge"))
        _add_teaching_flow_table(document, content.get("teaching_flow"), kp_names=kp_names)
        _add_lesson_detail_sections(document, content.get("session_plans"), kp_names=kp_names)
        _add_labeled_section(document, "课后安排", content.get("after_class_plan"), LESSON_AFTER_CLASS_LABELS)
        _add_list_section(document, "学情适配策略", content.get("learner_adjustments"))
        return _dump_document(document)

    def render_paper_result(self, paper_result, questions: list) -> bytes:
        """渲染试卷结果 DOCX。questions 需预先 enrich knowledge_point_name 字段。"""
        paper_json = _ensure_dict(paper_result.paper_json)
        document = _create_document()
        document.add_heading(strip_lesson_prefix(paper_result.title) or "试卷", level=0)
        _add_meta_paragraph(document, "试卷类型", labelize(paper_result.scene_type, SCENE_TYPE_LABELS, default="测评"))
        _add_meta_paragraph(document, "题目数量", paper_result.question_count)
        _add_distribution_section(document, "题型分布", paper_json.get("question_type_distribution"), QUESTION_TYPE_LABELS)
        _add_distribution_section(
            document,
            "难度分布",
            paper_json.get("difficulty_distribution"),
            {str(k): v for k, v in DIFFICULTY_LEVEL_LABELS.items()},
        )

        document.add_heading("题目明细", level=1)
        if not questions:
            document.add_paragraph("暂无题目。")
            return _dump_document(document)

        for question in questions:
            _add_question_section(document, question)
        return _dump_document(document)

    def render_homework_result(self, homework_result, questions: list, *, lesson_plan=None) -> bytes:
        """渲染课后作业结果 DOCX。questions 需预先 enrich knowledge_point_name 字段。"""
        content_json = _ensure_dict(homework_result.content_json)
        document = _create_document()
        document.add_heading(strip_lesson_prefix(homework_result.title) or "课后作业", level=0)
        scene_value = content_json.get("scene_type") or "homework"
        _add_meta_paragraph(document, "作业类型", labelize(scene_value, SCENE_TYPE_LABELS, default="课后作业"))
        if lesson_plan is not None:
            if lesson_plan.class_session_no is not None:
                _add_meta_paragraph(document, "课次", f"第 {lesson_plan.class_session_no} 讲")
            if lesson_plan.lesson_title:
                _add_meta_paragraph(document, "所属教案", strip_lesson_prefix(lesson_plan.lesson_title))
        _add_meta_paragraph(document, "题目数量", homework_result.question_count)
        _add_distribution_section(document, "题型分布", content_json.get("question_type_distribution"), QUESTION_TYPE_LABELS)
        _add_distribution_section(
            document,
            "难度分布",
            content_json.get("difficulty_distribution"),
            {str(k): v for k, v in DIFFICULTY_LEVEL_LABELS.items()},
        )

        document.add_heading("题目明细", level=1)
        if not questions:
            document.add_paragraph("暂无题目。")
            return _dump_document(document)

        for question in questions:
            _add_question_section(document, question)
        return _dump_document(document)


class DocumentExportService:
    """负责 DOCX 文件上传、文件对象落库与下载地址生成。"""

    def __init__(
        self,
        session: Session,
        storage_client: ObsStorageClient | None = None,
        render_service: DocxRenderService | None = None,
    ) -> None:
        self.session = session
        self.storage_client = storage_client or ObsStorageClient()
        self.render_service = render_service or DocxRenderService()
        self.settings = get_settings()

    def archive_docx(
        self,
        *,
        project_id: int,
        owner_user_id: int,
        biz_type: str,
        object_segments: tuple[str, ...],
        filename: str,
        content: bytes,
        metadata_json: dict[str, Any] | None,
        target,
    ) -> FileDownloadUrlResponse:
        """归档 DOCX 并回填业务对象导出文件。

        object_segments 末尾会追加模板版本号段（如 `tv2`），保证模板升级后 object_key 改变，
        旧 file_object 自然失效，前端再次触发 /export-docx 后落到新 key。
        """
        versioned_segments = (*object_segments, f"tv{DOCX_TEMPLATE_VERSION}")
        object_key = self.storage_client.build_object_key(*versioned_segments, filename=filename)
        try:
            self.storage_client.upload_bytes(object_key, content, content_type=DOCX_MIME_TYPE)
        except Exception as exc:  # noqa: BLE001
            raise AppException(BusinessErrorCode.FILE_UPLOAD_FAILED, "DOCX 文件上传失败", {"error": str(exc)}) from exc

        file_object = self._upsert_file_object(
            project_id=project_id,
            owner_user_id=owner_user_id,
            biz_type=biz_type,
            object_key=object_key,
            filename=filename,
            content=content,
            metadata_json=metadata_json,
        )
        target.export_file_id = file_object.id
        self.session.add(target)
        self.session.commit()

        try:
            signed_url = self.storage_client.create_download_signed_url(file_object.object_key)
        except Exception as exc:  # noqa: BLE001
            raise AppException(
                BusinessErrorCode.EXTERNAL_SERVICE_ERROR,
                "生成 DOCX 下载地址失败",
                {"file_object_id": file_object.id, "error": str(exc)},
            ) from exc

        return FileDownloadUrlResponse(
            file_object_id=file_object.id,
            bucket_name=file_object.bucket_name,
            object_key=file_object.object_key,
            signed_url=signed_url,
            expires_in_seconds=self.settings.obs_signed_url_expire_seconds,
            generated_at=DateTimeUtil.now_utc(),
        )

    def _upsert_file_object(
        self,
        *,
        project_id: int,
        owner_user_id: int,
        biz_type: str,
        object_key: str,
        filename: str,
        content: bytes,
        metadata_json: dict[str, Any] | None,
    ) -> FileObject:
        """按桶和对象路径复用或创建文件对象。"""
        bucket_name = self.storage_client.settings.obs_bucket
        statement = select(FileObject).where(FileObject.bucket_name == bucket_name, FileObject.object_key == object_key)
        file_object = self.session.scalar(statement)
        if file_object is None:
            file_object = FileObject(
                project_id=project_id,
                biz_type=biz_type,
                bucket_name=bucket_name,
                object_key=object_key,
                original_filename=filename,
                file_ext=".docx",
                mime_type=DOCX_MIME_TYPE,
                file_size=len(content),
                content_hash=hashlib.sha256(content).hexdigest(),
                source_type="system_export",
                upload_status="uploaded",
                uploaded_by=owner_user_id,
                metadata_json=metadata_json,
            )
        else:
            file_object.project_id = project_id
            file_object.biz_type = biz_type
            file_object.original_filename = filename
            file_object.file_ext = ".docx"
            file_object.mime_type = DOCX_MIME_TYPE
            file_object.file_size = len(content)
            file_object.content_hash = hashlib.sha256(content).hexdigest()
            file_object.source_type = "system_export"
            file_object.upload_status = "uploaded"
            file_object.uploaded_by = owner_user_id
            file_object.metadata_json = metadata_json
        self.session.add(file_object)
        self.session.flush()
        return file_object


def _create_document() -> Document:
    """创建基础 DOCX 文档。"""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    return document


def _dump_document(document: Document) -> bytes:
    """导出 DOCX 二进制内容。"""
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _ensure_dict(value: Any) -> dict[str, Any]:
    """确保结构化字段为字典。"""
    return value if isinstance(value, dict) else {}


def _safe_text(value: Any, fallback: str = "暂无") -> str:
    """转换安全展示文本。"""
    if value is None:
        return fallback
    if isinstance(value, str):
        return value.strip() or fallback
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _join_list(value: Any) -> str:
    """把列表或对象转换为多行文本。"""
    if isinstance(value, list):
        if not value:
            return "暂无"
        return "\n".join(_safe_text(item) for item in value)
    return _safe_text(value)


def _add_meta_paragraph(document: Document, label: str, value: Any) -> None:
    """追加键值段落。"""
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    paragraph.add_run(_safe_text(value))


def _add_plain_paragraph(document: Document, label: str, value: Any) -> None:
    """追加普通文本段落。"""
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    paragraph.add_run(_safe_text(value))


def _add_list_section(document: Document, title: str, values: Any) -> None:
    """追加列表章节。"""
    document.add_heading(title, level=1)
    if isinstance(values, list) and values:
        for item in values:
            document.add_paragraph(_safe_text(item), style="List Bullet")
        return
    document.add_paragraph(_safe_text(values))


def _add_labeled_section(
    document: Document,
    title: str,
    raw: Any,
    label_map: dict[str, str],
    *,
    level: int = 1,
) -> None:
    """按 label_map 的中文标签输出 dict 字段，未在 label_map 中的键会被跳过。"""
    pairs = iter_known_fields(raw, label_map)
    if not pairs:
        return
    document.add_heading(title, level=level)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for label, value in pairs:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = _join_list(value)


def _add_distribution_section(
    document: Document,
    title: str,
    raw: Any,
    label_map: dict[str, str],
    *,
    level: int = 1,
) -> None:
    """渲染题型/难度分布表，键转中文，未识别键跳过。"""
    if not isinstance(raw, dict) or not raw:
        return
    pairs: list[tuple[str, Any]] = []
    for key, value in raw.items():
        label = label_map.get(str(key))
        if label is None:
            continue
        pairs.append((label, value))
    if not pairs:
        return
    document.add_heading(title, level=level)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "类别"
    table.rows[0].cells[1].text = "数量"
    for label, value in pairs:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = _safe_text(value)


def _format_options(options_json: Any) -> str:
    """把题目选项渲染成 `A. xxx` 多行文本，dict/list 均兼容。"""
    if isinstance(options_json, dict) and options_json:
        # 选项字典通常以字母为键，按字母序输出便于教师阅读
        items = sorted(options_json.items(), key=lambda kv: str(kv[0]))
        return "\n".join(f"{key}. {_safe_text(value)}" for key, value in items)
    if isinstance(options_json, list) and options_json:
        lines: list[str] = []
        for index, item in enumerate(options_json):
            if isinstance(item, dict) and "label" in item and "text" in item:
                lines.append(f"{item.get('label')}. {_safe_text(item.get('text'))}")
            else:
                lines.append(f"{chr(ord('A') + index)}. {_safe_text(item)}")
        return "\n".join(lines)
    return ""


def _format_knowledge_points(refs: Any, names: dict[int, str]) -> str:
    """把 knowledge_point_refs（int 列表）转成知识点名称的换行文本。"""
    if not isinstance(refs, list) or not refs:
        return ""
    labels: list[str] = []
    for ref in refs:
        try:
            ref_int = int(ref)
        except (TypeError, ValueError):
            labels.append(_safe_text(ref))
            continue
        name = names.get(ref_int)
        labels.append(name if name else f"知识点{ref_int}")
    return "\n".join(labels)


def _add_question_section(document: Document, question: Any) -> None:
    """渲染单题：题号 → 题型/难度/分值 → 题干 → 选项 → 答案 → 解析 → 知识点。"""
    document.add_heading(f"第 {_safe_text(question.question_no)} 题", level=2)
    _add_meta_paragraph(document, "题型", labelize(question.question_type, QUESTION_TYPE_LABELS, default="未分类"))
    _add_meta_paragraph(
        document,
        "难度",
        labelize(question.difficulty_level, DIFFICULTY_LEVEL_LABELS, default="—"),
    )
    if question.score_value is not None:
        _add_meta_paragraph(document, "分值", question.score_value)
    _add_plain_paragraph(document, "题干", question.stem_text)
    options_text = _format_options(getattr(question, "options_json", None))
    if options_text:
        paragraph = document.add_paragraph()
        paragraph.add_run("选项：").bold = True
        lines = options_text.splitlines()
        paragraph.add_run(lines[0])
        for line in lines[1:]:
            paragraph.add_run().add_break(WD_BREAK.LINE)
            paragraph.add_run(line)
    _add_plain_paragraph(document, "答案", question.answer_text)
    if question.analysis_text:
        _add_plain_paragraph(document, "解析", question.analysis_text)
    knowledge_point_name = getattr(question, "knowledge_point_name", None)
    if knowledge_point_name:
        _add_meta_paragraph(document, "知识点", knowledge_point_name)


def _add_lesson_session_table(document: Document, sessions: Any) -> None:
    """追加课次安排表。"""
    document.add_heading("课次安排", level=1)
    if not isinstance(sessions, list) or not sessions:
        document.add_paragraph("暂无课次安排。")
        return
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["课次", "标题", "时长", "目标", "重点", "活动", "课后任务"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for session in sessions:
        item = _ensure_dict(session)
        cells = table.add_row().cells
        cells[0].text = _safe_text(item.get("session_no"))
        cells[1].text = _safe_text(item.get("title"))
        cells[2].text = _safe_text(item.get("duration_minutes"))
        cells[3].text = _join_list(item.get("objectives"))
        cells[4].text = _join_list(item.get("key_points"))
        cells[5].text = _join_list(item.get("activities"))
        cells[6].text = _join_list(item.get("homework"))


def _add_teaching_flow_table(document: Document, steps: Any, *, kp_names: dict[int, str]) -> None:
    """追加教学流程表，知识点列用名称替代 ID。"""
    document.add_heading("教学流程", level=1)
    if not isinstance(steps, list) or not steps:
        document.add_paragraph("暂无教学流程。")
        return
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["序号", "环节", "时长", "教师动作", "学生活动", "知识点"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for step in steps:
        item = _ensure_dict(step)
        cells = table.add_row().cells
        cells[0].text = _safe_text(item.get("step_no"))
        cells[1].text = _safe_text(item.get("stage_name"))
        cells[2].text = _safe_text(item.get("duration_minutes"))
        cells[3].text = _join_list(item.get("teacher_actions"))
        cells[4].text = _join_list(item.get("student_activities"))
        cells[5].text = _format_knowledge_points(item.get("knowledge_point_refs"), kp_names) or "—"


def _add_lesson_detail_sections(document: Document, session_plans: Any, *, kp_names: dict[int, str]) -> None:
    """追加课次讲解安排。"""
    document.add_heading("课次讲解安排", level=1)
    if not isinstance(session_plans, list) or not session_plans:
        document.add_paragraph("暂无课次讲解安排。")
        return
    for session_plan in session_plans:
        item = _ensure_dict(session_plan)
        document.add_heading(strip_lesson_prefix(_safe_text(item.get("title"), "课次安排")), level=2)
        _add_list_like_paragraph(document, "课次目标", item.get("objectives"))
        _add_list_like_paragraph(document, "教学重点", item.get("teaching_focus"))
        _add_list_like_paragraph(document, "课后任务", item.get("homework"))
        knowledge_points_text = _format_knowledge_points(item.get("knowledge_point_refs"), kp_names)
        if knowledge_points_text:
            _add_list_like_paragraph(document, "涉及知识点", knowledge_points_text.split("\n"))
        _add_teaching_step_list(document, item.get("teaching_steps"))


def _add_list_like_paragraph(document: Document, label: str, values: Any) -> None:
    """追加带换行的列表型段落。"""
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    text = _join_list(values)
    lines = text.splitlines() or ["暂无"]
    paragraph.add_run(lines[0])
    for line in lines[1:]:
        paragraph.add_run().add_break(WD_BREAK.LINE)
        paragraph.add_run(line)


def _add_teaching_step_list(document: Document, steps: Any) -> None:
    """追加课次内教学步骤。"""
    if not isinstance(steps, list) or not steps:
        return
    for step in steps:
        item = _ensure_dict(step)
        text = (
            f"{_safe_text(item.get('step_no'))}. {_safe_text(item.get('stage_name'))}："
            f"教师动作 {_join_list(item.get('teacher_actions'))}；"
            f"学生活动 {_join_list(item.get('student_activities'))}"
        )
        document.add_paragraph(text, style="List Number")
